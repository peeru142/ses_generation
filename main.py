from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
import win32com.client
import time
import os
import shutil
import re
import datetime
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import traceback

app = FastAPI()

# CORS configuration for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://localhost:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# SAP Login Details
SAP_CONNECTION = "S4HANA"
SAP_CLIENT = "300"
SAP_USER = "U10009674"
SAP_PASS = "Lahore@123"
SAP_LANG = "EN"

# Template path
TEMPLATE_PATH = r"SESinfo.docx"

# Output directory for generated files
OUTPUT_DIR = os.path.join(os.path.dirname(TEMPLATE_PATH), "SES_Generated")
os.makedirs(OUTPUT_DIR, exist_ok=True)

class PORequest(BaseModel):
    po_number: int
    distribution_name: str

class SAPExtractor:
    def __init__(self):
        self.session = None
        self.connection = None
        
    def connect_to_sap(self):
        """Connect to SAP GUI"""
        try:
            SapGuiAuto = win32com.client.GetObject("SAPGUI")
            application = SapGuiAuto.GetScriptingEngine
            self.connection = application.OpenConnection(SAP_CONNECTION, True)
            self.session = self.connection.Children(0)
            return True
        except Exception as e:
            raise Exception(f"Failed to connect to SAP: {str(e)}")
    
    def login(self):
        """Login to SAP"""
        try:
            self.session.findById("wnd[0]/usr/txtRSYST-MANDT").text = SAP_CLIENT
            self.session.findById("wnd[0]/usr/txtRSYST-BNAME").text = SAP_USER
            self.session.findById("wnd[0]/usr/pwdRSYST-BCODE").text = SAP_PASS
            self.session.findById("wnd[0]/usr/txtRSYST-LANGU").text = SAP_LANG
            self.session.findById("wnd[0]").sendVKey(0)
            time.sleep(2)
        except Exception as e:
            raise Exception(f"Login failed: {str(e)}")
    
    def open_purchase_order(self, po_number):
        """Navigate to ME23N and open PO"""
        try:
            self.session.findById("wnd[0]/tbar[0]/okcd").text = "ME23N"
            self.session.findById("wnd[0]").sendVKey(0)
            time.sleep(2)
            
            self.session.findById("wnd[0]/tbar[1]/btn[17]").press()
            time.sleep(1)
            
            po_input_field = self.session.findById("wnd[1]/usr/subSUB0:SAPLMEGUI:0003/ctxtMEPO_SELECT-EBELN")
            po_input_field.text = ""
            po_input_field.text = str(po_number)
            po_input_field.caretPosition = len(str(po_number))
            
            self.session.findById("wnd[1]").sendVKey(0)
            time.sleep(2)
        except Exception as e:
            raise Exception(f"Failed to open PO: {str(e)}")
    
    def clean_vendor_name(self, raw_vendor_text):
        """Remove vendor ID numbers and clean the vendor name"""
        if not raw_vendor_text:
            return ""
        cleaned = raw_vendor_text.strip()
        cleaned = re.sub(r'^\d+\s+', '', cleaned)
        return cleaned
    
    def extract_vendor_name(self):
        """Extract vendor name from PO"""
        vendor_name = ""
        
        # Method 1: Try the primary field
        try:
            vendor_name_field = self.session.findById("wnd[0]/usr/subSCREEN_3000_RESIZING_AREA:SAPLBUS_LOCATOR:2036/subSCREEN_1010_RIGHT_AREA:SAPLBUPA_DIALOG_JOEL:1000/ssubSCREEN_1000_WORKAREA_AREA:SAPLBUPA_DIALOG_JOEL:1100/ssubSCREEN_1100_MAIN_AREA:SAPLBUPA_DIALOG_JOEL:1101/tabsGS_SCREEN_1100_TABSTRIP/tabpSCREEN_1100_TAB_01/ssubSCREEN_1100_TABSTRIP_AREA:SAPLBUSS:0028/ssubGENSUB:SAPLBUSS:7016/subA02P02:SAPLBUD0:1200/txtBUT000-NAME_ORG2")
            raw_vendor_name = vendor_name_field.text
            if raw_vendor_name.strip():
                vendor_name = self.clean_vendor_name(raw_vendor_name)
                return vendor_name
        except:
            pass
        
        # Method 2: Try alternative vendor fields
        alternative_vendor_fields = [
            "wnd[0]/usr/subSUB0:SAPLMEGUI:0019/subSUB0:SAPLMEGUI:0030/subSUB1:SAPLMEGUI:1105/ctxtMEPO_TOPLINE-SUPERFIELD",
            "wnd[0]/usr/subSUB0:SAPLMEGUI:0019/subSUB1:SAPLMEGUI:1100/txtLFA1-NAME1",
            "wnd[0]/usr/subSUB0:SAPLMEGUI:0019/subSUB2:SAPLMEGUI:1200/txtLFA1-NAME1"
        ]
        
        for field_id in alternative_vendor_fields:
            try:
                vendor_field = self.session.findById(field_id)
                raw_vendor_name = vendor_field.text
                if raw_vendor_name.strip():
                    vendor_name = self.clean_vendor_name(raw_vendor_name)
                    return vendor_name
            except:
                continue
        
        return ""
    
    def clean_amount(self, raw_amount):
        """Return cleaned numeric string"""
        if not raw_amount:
            return ""
        m = re.search(r'[-+]?[0-9][0-9,\.]*', raw_amount)
        if not m:
            return ""
        token = m.group(0).replace(",", "")
        if "." in token:
            int_part, dec_part = token.split(".", 1)
            if re.fullmatch(r'0+', dec_part):
                token = int_part
        return token
    
    def extract_service_lines(self):
        """Extract service lines with amounts"""
        services_found = []
        
        try:
            service_table_path = (
                "wnd[0]/usr/subSUB0:SAPLMEGUI:0019/"
                "subSUB2:SAPLMEVIEWS:1100/subSUB2:SAPLMEVIEWS:1200/"
                "subSUB1:SAPLMEGUI:1211/tblSAPLMEGUITC_1211"
            )
            
            service_table = self.session.findById(service_table_path)
            scrollbar = service_table.verticalScrollbar
            visible_rows = service_table.VisibleRowCount
            seen_keys = set()
            
            def extract_row_data(rel_row):
                line_number = ""
                for line_col in [0, 1, 10, 2, 3]:
                    try:
                        fid = f"{service_table_path}/txtMEPO1211-EBELP[{line_col},{rel_row}]"
                        val = self.session.findById(fid).text.strip()
                        if val:
                            line_number = val
                            break
                    except:
                        continue
                
                service_name = ""
                for service_col in [5, 4, 6, 3, 7, 8]:
                    try:
                        fid = f"{service_table_path}/txtMEPO1211-TXZ01[{service_col},{rel_row}]"
                        val = self.session.findById(fid).text.strip()
                        if val:
                            service_name = val
                            break
                    except:
                        continue
                
                invoice_amount_raw = ""
                candidate_fields = ["WRBTR", "NETWR", "NETPR", "KBETR", "BRTWR", "DMBTR"]
                candidate_cols = [15, 14, 16, 12, 13, 17, 18, 19, 20, 10, 11, 9]
                
                for fld in candidate_fields:
                    for col in candidate_cols:
                        try:
                            fid = f"{service_table_path}/txtMEPO1211-{fld}[{col},{rel_row}]"
                            val = self.session.findById(fid).text.strip()
                            if val and re.search(r'\d', val):
                                invoice_amount_raw = val
                                break
                        except:
                            continue
                    if invoice_amount_raw:
                        break
                
                invoice_amount_clean = self.clean_amount(invoice_amount_raw)
                return line_number, service_name, invoice_amount_raw, invoice_amount_clean
            
            scrollbar.position = 0
            time.sleep(0.3)
            
            scroll_attempts = 0
            max_scroll_attempts = 60
            
            while scroll_attempts < max_scroll_attempts:
                for rel_row in range(visible_rows):
                    try:
                        line_number, service_name, invoice_raw, invoice_clean = extract_row_data(rel_row)
                        
                        key = (line_number, service_name)
                        if line_number and key not in seen_keys:
                            services_found.append({
                                "line": line_number,
                                "service": service_name,
                                "amount_raw": invoice_raw,
                                "amount": invoice_clean
                            })
                            seen_keys.add(key)
                    except:
                        continue
                
                current_pos = scrollbar.position
                next_position = min(current_pos + visible_rows - 1, scrollbar.maximum)
                if next_position <= current_pos:
                    break
                
                scrollbar.position = next_position
                time.sleep(0.25)
                scroll_attempts += 1
            
            services_found.sort(key=lambda x: int(x['line']) if x['line'].isdigit() else 9999)
            
        except Exception as e:
            print(f"Error extracting services: {e}")
        
        return services_found
    
    def disconnect(self):
        """Close SAP connection"""
        try:
            if self.connection:
                self.connection.CloseSession(self.session.Id)
        except:
            pass

def add_bullet(paragraph, text):
    """
    Turn a paragraph into a bulleted item with given text.
    """
    # clear any existing text
    paragraph.clear()
    run = paragraph.add_run(text)

    # Apply bullet numbering (list style) using XML
    p = paragraph._p  # lxml element
    pPr = p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')

    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), "0")
    numPr.append(ilvl)

    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), "1")  # bullet list id
    numPr.append(numId)

    pPr.append(numPr)

def fill_ses_template(template_path, po_number, vendor_name, service_items, distribution_name):
    """Fill SES template with extracted data"""
    try:
        base_name = os.path.basename(template_path)
        name, ext = os.path.splitext(base_name)

        # Always save to OUTPUT_DIR
        new_file_path = os.path.join(OUTPUT_DIR, f"{name}_PO_{po_number}{ext}")
        shutil.copyfile(template_path, new_file_path)
        
        doc = Document(new_file_path)
        current_date = datetime.datetime.now().strftime("%d.%m.%Y")

        # Build amount lines (formatted with commas)
        amount_lines = []
        for s in service_items:
            line_num = s.get("line", "").strip()
            amt = s.get("amount", "").strip()
            
            if amt and amt.isdigit():
                amt_formatted = "{:,}".format(int(amt))
            else:
                amt_formatted = amt

            if amt_formatted:
                line_text = f"Amount Rs: {amt_formatted} - line item# {line_num}"
            else:
                line_text = f"Amount Rs: - line item# {line_num}"

            amount_lines.append(line_text)

        # --- Replace table fields (Date and Distribution) ---
        date_replaced = False
        distribution_replaced = False
        for table in doc.tables:
            for row in table.rows:
                cells_text = [cell.text for cell in row.cells]
                if not date_replaced and any("Date:" in t for t in cells_text):
                    for idx, cell in enumerate(row.cells):
                        if "Date:" in cell.text and idx + 1 < len(row.cells):
                            for p in row.cells[idx + 1].paragraphs:
                                p.text = current_date
                            date_replaced = True
                            break
                if not distribution_replaced and any("Distribution:" in t for t in cells_text):
                    for idx, cell in enumerate(row.cells):
                        if "Distribution:" in cell.text and idx + 1 < len(row.cells):
                            for p in row.cells[idx + 1].paragraphs:
                                p.text = distribution_name
                            distribution_replaced = True
                            break
                if date_replaced and distribution_replaced:
                    break
            if date_replaced and distribution_replaced:
                break

        # --- Replace placeholders in text ---
        replaced_amounts = False
        paragraph_to_remove = None
        
        for p in doc.paragraphs:
            txt = p.text

            if "PO#" in txt:
                txt = re.sub(r'PO#\s*\d*', f'PO# {po_number}', txt)

            if "in favor of" in txt and vendor_name:
                parts = txt.split("in favor of")
                if len(parts) >= 2:
                    before = parts[0] + "in favor of"
                    after = parts[1]
                    if "for taking" in after:
                        after_parts = after.split("for taking", 1)
                        txt = f'{before} "{vendor_name}" for taking{after_parts[1]}'
                    else:
                        txt = f'{before} "{vendor_name}"'

            if "for taking the services of" in txt:
                services = [s["service"] for s in service_items if s.get("service")]
                if services:
                    all_services = ", ".join(services)
                    parts = txt.split("for taking the services of")
                    txt = f'{parts[0]}for taking the services of "{all_services}"'

            # If amount placeholder exists → replace with bullet list
            if "Amount Rs:" in txt and "line item" in txt:
                if amount_lines:
                    paragraph_to_remove = p  # Mark for removal
                    # Insert bullet points BEFORE this paragraph
                    for line in amount_lines:
                        new_p = p.insert_paragraph_before("")
                        add_bullet(new_p, line)
                    replaced_amounts = True
                continue

            # For all other cases → update normally
            if txt != p.text:
                p.text = txt

        # Remove the original placeholder paragraph (which is now empty with a bullet)
        if paragraph_to_remove is not None:
            paragraph_to_remove._element.getparent().remove(paragraph_to_remove._element)

        # If no placeholder existed → append bullet list at end
        if not replaced_amounts and amount_lines:
            for line in amount_lines:
                new_p = doc.add_paragraph("")
                add_bullet(new_p, line)

        doc.save(new_file_path)
        print(f"Document saved: {new_file_path}")
        return new_file_path

    except Exception as e:
        print(f"Error filling template: {e}")
        traceback.print_exc()
        return None

@app.get("/")
def read_root():
    return {"message": "SES Generator API"}

@app.post("/connect-sap")
async def connect_sap():
    """Test SAP connection endpoint"""
    try:
        return {
            "success": True,
            "message": "SAP will connect automatically when generating SES"
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")

@app.post("/generate-ses")
async def generate_ses(request: PORequest):
    """Generate SES document from PO number"""
    extractor = SAPExtractor()
    
    try:
        print(f"\nProcessing PO: {request.po_number}")
        print(f"Distribution Name: {request.distribution_name}")
        
        extractor.connect_to_sap()
        extractor.login()
        extractor.open_purchase_order(request.po_number)
        
        vendor_name = extractor.extract_vendor_name()
        service_lines = extractor.extract_service_lines()
        
        if not service_lines:
            raise HTTPException(status_code=404, detail="No service lines found in PO")
        
        output_file = fill_ses_template(
            TEMPLATE_PATH, 
            request.po_number, 
            vendor_name, 
            service_lines,
            request.distribution_name
        )
        
        if not output_file or not os.path.exists(output_file):
            raise HTTPException(status_code=500, detail="Failed to generate document")
        
        return {
            "success": True,
            "po_number": request.po_number,
            "vendor_name": vendor_name if vendor_name else "[NOT FOUND]",
            "service_count": len(service_lines),
            "file_path": output_file
        }
        
    except Exception as e:
        print(f"ERROR: {str(e)}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")
    
    finally:
        extractor.disconnect()

@app.get("/download/{po_number}")
async def download_file(po_number: int):
    base_name = os.path.basename(TEMPLATE_PATH)
    name, ext = os.path.splitext(base_name)
    file_path = os.path.join(OUTPUT_DIR, f"{name}_PO_{po_number}{ext}")
    
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail=f"File not found: {file_path}")
    
    return FileResponse(
        path=file_path,
        filename=os.path.basename(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)